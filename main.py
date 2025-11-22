import argparse
from ultralytics import YOLO
from ultralytics.utils import LOGGER
from deepface import DeepFace
import logging
import io
import cv2
import os
import time
from docx import Document
from docx.shared import Inches
import re


doc = Document()
save_dir = "output_images"
save_classify = "output_classification"
save_documents = "Documents"



def get_highest(dic):
    
    biggest = 0
    answer = None
    for i,j in dic.items():
        
        if j>biggest:
            
            biggest = j 
            answer = i

    return answer

def classify_identify(image,classify,identify):

    
    
    model_class = ""
    if classify != None:
        model_class = classify
    else:
        model_class = "yolo11x.pt" # default

    model_ident = ""
    if identify != None:
        model_ident = identify
    else:
        model_ident = "yolo11x-cls.pt" # default

    print(model_class,model_ident)

    # Classify
    model = YOLO(model_ident)
    results = model(image)
    

    # Potential items
    model_1 = YOLO(model_class)
    results_1 = model_1(image)
    

    return (results,results_1)

def output_to_doc(output):

    # save
    for i in output[0]:
        img = i.plot()
        save_path = os.path.join(save_dir,f"output{int(time.time()*1000)}.jpg")

        cv2.imwrite(save_path,img)
        doc.add_heading("Identification Results")
        doc.add_picture(save_path, width= Inches(4), height=None)
        

    for i in output[1]:
        img = i.plot()
        save_path = os.path.join(save_classify,f"output{int(time.time()*1000)}.jpg")
        doc.add_heading("Classification Results")
        
        cv2.imwrite(save_path,img)
        doc.add_picture(save_path, width=Inches(4), height=None)

    return 

def analysis_output(image):

    objs = DeepFace.analyze(img_path = image , actions = ['age', 'gender', 'race', 'emotion'])
    doc.add_heading("Analysis Results")
    doc.add_paragraph(f"{objs[0]["age"]},{get_highest(objs[0]["gender"])},{get_highest(objs[0]["race"])},{get_highest(objs[0]["emotion"])}")

    return 



def main():

    # Create a string buffer
    log_buffer = io.StringIO()

    # Create a logging handler that writes into the buffer
    handler = logging.StreamHandler(log_buffer)
    handler.setLevel(logging.INFO)
    LOGGER.addHandler(handler)

    save_dir = "output_images"
    save_classify = "output_classification"

    parser = argparse.ArgumentParser(
                    prog='Facial analysis and image classification',
                    description='Image classification')
    
    parser.add_argument("--mode",choices = ["both","analysis","classification"],required = True)
    parser.add_argument("--classification",type = str,required = False)
    parser.add_argument("--identification",type = str ,required = False)
    parser.add_argument("--img",type = str, required = True)

    args = parser.parse_args()

    mode = args.mode
    image = args.img
    model_classification = args.classification
    model_identifcation = args.identification

    if mode == "both":

        output = classify_identify(image,model_classification,model_identifcation)
        output_to_doc(output)
        analysis_output(image)

    elif mode == "analysis":

        analysis_output(image)

    elif mode == "classification":
        
        output = classify_identify(image,model_classification,model_identifcation)
        output_to_doc(output)


    handler.flush()
    captured = log_buffer.getvalue()
    tidied = ""
    for line in captured.splitlines():
        
        text = re.sub(r'(image \d+/\d+ )(.+?[/\\])([^/\\:]+\.jpg:)', r'\1\3', line)
        tidied = tidied + text + "\n"



    doc.add_heading("Logger Results")
    doc.add_paragraph(tidied)
  

    return 




if __name__ == "__main__":
    
    doc.add_heading("Analysis Report", level=1)
    main()
    

    save_path = os.path.join(save_documents,f"detections{int(time.time()*1000)}.docx")
    doc.save(save_path)
       